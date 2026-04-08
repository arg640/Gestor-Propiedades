import re
import io
from datetime import datetime
from zoneinfo import ZoneInfo
from copy import deepcopy

from docx import Document
from num2words import num2words

MEXICO_TZ = ZoneInfo("America/Mexico_City")

MESES_ES = {
    1: "Enero", 2: "Febrero", 3: "Marzo", 4: "Abril",
    5: "Mayo", 6: "Junio", 7: "Julio", 8: "Agosto",
    9: "Septiembre", 10: "Octubre", 11: "Noviembre", 12: "Diciembre"
}

def monto_a_letras(monto_str: str) -> str:
    """
    Convierte '$10,000.00' → 'DIEZ MIL PESOS 00/100 M.N.'
    """
    limpio = monto_str.replace("$", "").replace(" MXN", "").replace(",", "").strip()
    partes = limpio.split(".")
    entero = int(partes[0])
    centavos = int(partes[1].ljust(2, "0")[:2]) if len(partes) > 1 else 0
    texto = num2words(entero, lang="es").upper()
    return f"{texto} PESOS {centavos:02d}/100 M.N."

def _merge_runs(paragraph):
    """
    Fusiona runs consecutivos con el mismo atributo bold para que cada
    placeholder {{...}} quede en un solo run y se pueda reemplazar fácilmente.
    """
    runs = paragraph.runs
    if len(runs) < 2:
        return
    i = 0
    while i < len(paragraph.runs) - 1:
        curr = paragraph.runs[i]
        nxt = paragraph.runs[i + 1]
        if curr.bold == nxt.bold:
            curr.text += nxt.text
            nxt._element.getparent().remove(nxt._element)
        else:
            i += 1

def _replace_in_paragraph(paragraph, etiquetas: dict):
    """
    Reemplaza todas las etiquetas {{clave}} en el párrafo.
    Fusiona runs primero para manejar placeholders fragmentados,
    luego hace el reemplazo preservando bold/italic del run original.
    """
    _merge_runs(paragraph)
    for run in paragraph.runs:
        for clave, valor in etiquetas.items():
            placeholder = f"{{{{{clave}}}}}"
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, str(valor))

def generar_recibo(datos: dict, ruta_plantilla: str) -> bytes:
    """
    Toma los datos extraídos del contrato y genera el recibo como
    bytes de un archivo .docx, listo para descargarse.

    Etiquetas soportadas en la plantilla:
      {{nombre_inquilino}}, {{nombre_arrendador}} / {{nombre_administrador}},
      {{monto_num}}, {{monto_letra}}, {{mes_renta}},
      {{num_finca}}, {{direccion_coto}}, {{fecha_recibo}}
    """
    hoy = datetime.now(tz=MEXICO_TZ)   # horario Ciudad de México
    mes_anio = f"{MESES_ES[hoy.month]} {hoy.year}"
    fecha_recibo = f"{hoy.day:02d} de {MESES_ES[hoy.month]} de {hoy.year}"

    monto_str = datos.get("monto_renta", "$0.00 MXN")
    monto_num = monto_str.replace(" MXN", "")
    monto_letra = monto_a_letras(monto_str)

    nombre_admin = datos.get("nombre_arrendador") or datos.get("nombre_administrador", "")

    etiquetas = {
        "nombre_inquilino":    datos.get("nombre_inquilino", ""),
        "nombre_arrendador":   nombre_admin,
        "nombre_administrador": nombre_admin,
        "domicilio_dueno":     datos.get("domicilio_dueno", ""),
        "monto_num":           monto_num,
        "monto_letra":         monto_letra,
        "mes_renta":           mes_anio,
        "num_finca":           datos.get("num_finca", ""),
        "direccion_coto":      datos.get("direccion_coto", ""),
        "fecha_recibo":        fecha_recibo,
    }

    print("=" * 60)
    print("GENERANDO RECIBO CON ETIQUETAS:")
    for k, v in etiquetas.items():
        print(f"  {{{{{k}}}}} -> {v}")
    print("=" * 60)

    doc = Document(ruta_plantilla)

    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, etiquetas)

    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for paragraph in celda.paragraphs:
                    _replace_in_paragraph(paragraph, etiquetas)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
